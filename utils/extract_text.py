import os
import io
import fitz  # PyMuPDF
import docx
from pdfminer.high_level import extract_text as pdfminer_extract_text
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage


def extract_text_from_pdf_pymupdf(file_path):
    """Extract text from PDF using PyMuPDF"""
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()  # Close the document to free resources
        return text
    except Exception as e:
        print(f"PyMuPDF extraction failed: {e}")
        return None


def extract_text_from_pdf_pdfminer(file_path):
    """Extract text from PDF using pdfminer as a backup method"""
    try:
        return pdfminer_extract_text(file_path)
    except Exception as e:
        print(f"PDFMiner direct extraction failed: {e}")
        
        # Try alternative method with pdfminer
        try:
            resource_manager = PDFResourceManager()
            fake_file_handle = io.StringIO()
            converter = TextConverter(resource_manager, fake_file_handle, laparams=LAParams())
            page_interpreter = PDFPageInterpreter(resource_manager, converter)
            
            with open(file_path, 'rb') as fh:
                for page in PDFPage.get_pages(fh, caching=True, check_extractable=True):
                    page_interpreter.process_page(page)
                
            text = fake_file_handle.getvalue()
            
            # Clean up resources
            converter.close()
            fake_file_handle.close()
            
            return text
        except Exception as e:
            print(f"PDFMiner alternative method failed: {e}")
            return None


def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
                    
        return '\n'.join(full_text)
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return None


def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read()
        except UnicodeDecodeError:
            continue  # Try next encoding
        except Exception as e:
            print(f"TXT extraction failed with {encoding}: {e}")
    
    print("All encoding attempts failed for TXT file")
    return None


def extract_text_from_file(file_path):
    """Extract text from file based on extension"""
    if not os.path.exists(file_path):
        return f"Error: File '{file_path}' does not exist."
    
    try:
        _, file_extension = os.path.splitext(file_path)
        file_extension = file_extension.lower()
        
        if file_extension == '.pdf':
            # Try PyMuPDF first, fallback to pdfminer
            text = extract_text_from_pdf_pymupdf(file_path)
            if not text:
                text = extract_text_from_pdf_pdfminer(file_path)
            return text or "PDF text extraction failed."
            
        elif file_extension == '.docx':
            text = extract_text_from_docx(file_path)
            return text or "DOCX text extraction failed."
            
        elif file_extension == '.doc':
            # Add warning for .doc files since python-docx doesn't fully support them
            print("Warning: .doc files may not be fully supported. Consider converting to .docx")
            text = extract_text_from_docx(file_path)
            return text or "DOC text extraction failed. Consider converting to DOCX format."
            
        elif file_extension == '.txt':
            text = extract_text_from_txt(file_path)
            return text or "TXT text extraction failed."
            
        else:
            return f"Unsupported file format: {file_extension}"
    except Exception as e:
        return f"An unexpected error occurred: {str(e)}"


# Example usage
if __name__ == "__main__":
    file_path = "example.pdf"  # Replace with your file path
    text = extract_text_from_file(file_path)
    print(text[:500])  # Print first 500 characters as a preview
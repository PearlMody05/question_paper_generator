import os
import google.generativeai as genai
from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from werkzeug.utils import secure_filename
from utils.extract_text import extract_text_from_file
import io
import tempfile
from docx import Document
from dotenv import load_dotenv
load_dotenv()

# Configure Flask app
app = Flask(__name__)
app.secret_key = "Imgonnagetuback"  # For flash messages
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload size
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'doc', 'docx', 'txt'}

# Create uploads folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Configure Gemini API
# Replace with your actual API key when deploying
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

# Initialize Gemini model
model = genai.GenerativeModel('gemini-1.5-pro')

def allowed_file(filename):
    """Check if the file extension is allowed"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.route('/')
def index():
    """Render the main page with the upload form"""
    return render_template('index.html')

@app.route('/predict', methods=['POST'])
def predict():
    """Process uploaded files and generate question paper predictions"""
    if request.method == 'POST':
        # Check if all required fields are provided
        course_outcomes = request.form.get('course_outcomes', '')
        total_marks = request.form.get('total_marks', '')
        
        if not course_outcomes or not total_marks:
            flash('Course Outcomes and Total Marks are required.')
            return redirect(url_for('index'))
        
        # Process file uploads
        previous_papers_text = ""
        syllabus_text = ""
        notes_text = ""
        
        # Process Previous Papers
        if 'previous_papers' not in request.files:
            flash('Previous Papers are required.')
            return redirect(url_for('index'))
        
        previous_papers = request.files.getlist('previous_papers')
        if not previous_papers or previous_papers[0].filename == '':
            flash('Previous Papers are required.')
            return redirect(url_for('index'))
        
        for file in previous_papers:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                extracted_text = extract_text_from_file(filepath)
                previous_papers_text += f"\n--- BEGIN {filename} ---\n"
                previous_papers_text += extracted_text
                previous_papers_text += f"\n--- END {filename} ---\n"
                # Clean up file
                os.remove(filepath)
        
        # Process Syllabus
        if 'syllabus' not in request.files:
            flash('Syllabus is required.')
            return redirect(url_for('index'))
        
        syllabus_file = request.files['syllabus']
        if syllabus_file and syllabus_file.filename != '' and allowed_file(syllabus_file.filename):
            filename = secure_filename(syllabus_file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            syllabus_file.save(filepath)
            syllabus_text = extract_text_from_file(filepath)
            # Clean up file
            os.remove(filepath)
        else:
            flash('Syllabus file is required.')
            return redirect(url_for('index'))
        
        # Process Notes (optional)
        if 'notes' in request.files:
            notes_file = request.files['notes']
            if notes_file and notes_file.filename != '' and allowed_file(notes_file.filename):
                filename = secure_filename(notes_file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                notes_file.save(filepath)
                notes_text = extract_text_from_file(filepath)
                # Clean up file
                os.remove(filepath)
        
        # Create prompt for Gemini
        prompt = f"""You are an educational question paper prediction assistant designed for teachers.
You are given:
* Previous year question papers (text): {previous_papers_text}
* Syllabus content (unit-wise topics): {syllabus_text}
* Optional class notes: {notes_text}
* Course Outcomes (COs): {course_outcomes}
* Total Marks of the exam: {total_marks}

Your task is to produce three sections:
📌 SECTION 1: Previously Asked Repeated Questions
* List questions that have appeared multiple times in the previous papers
* Bold them or mark them clearly
* Include unique year if possible

📌 SECTION 2: Most Probable Questions
* Based on syllabus topics, trends in previous papers, and content from notes
* Use Bloom's taxonomy levels (e.g., Understand, Apply, Analyze)
* Align with Course Outcomes

📌 SECTION 3: Generated Question Paper
* Well-structured exam-style question paper
* Questions should span Bloom's levels and multiple Course Outcomes
* Distribute questions to match the given total marks (e.g., 2M, 5M, 10M format)
* Highlight if any questions overlap with repeated ones from Section 1

Be smart and predictive — use topic weightage, historical patterns, and question variety. Ensure formatting is clear.
Format your response with Markdown for better readability.
"""
        
        try:
            # Get response from Gemini
            gemini_response = model.generate_content(prompt)
            response_text = gemini_response.text
            
            # Extract the three sections from the response
            sections = {}
            
            # Simple parsing - improve as needed
            if "SECTION 1: Previously Asked Repeated Questions" in response_text:
                section1_start = response_text.find("SECTION 1: Previously Asked Repeated Questions")
                section2_start = response_text.find("SECTION 2: Most Probable Questions")
                sections["repeated_questions"] = response_text[section1_start:section2_start].strip()
            else:
                sections["repeated_questions"] = "No repeated questions identified."
            
            if "SECTION 2: Most Probable Questions" in response_text:
                section2_start = response_text.find("SECTION 2: Most Probable Questions")
                section3_start = response_text.find("SECTION 3: Generated Question Paper")
                sections["probable_questions"] = response_text[section2_start:section3_start].strip()
            else:
                sections["probable_questions"] = "No probable questions identified."
            
            if "SECTION 3: Generated Question Paper" in response_text:
                section3_start = response_text.find("SECTION 3: Generated Question Paper")
                sections["generated_paper"] = response_text[section3_start:].strip()
            else:
                sections["generated_paper"] = "No question paper generated."
            
            return render_template('result.html', 
                                   sections=sections,
                                   course_outcomes=course_outcomes,
                                   total_marks=total_marks,
                                   full_response=response_text)
        
        except Exception as e:
            flash(f"Error generating question paper: {str(e)}")
            return redirect(url_for('index'))

@app.route('/download/<format>')
def download(format):
    """Download the generated question paper in specified format"""
    full_response = request.args.get('full_response', '')
    
    if format == 'txt':
        buffer = io.BytesIO()
        buffer.write(full_response.encode('utf-8'))
        buffer.seek(0)
        return send_file(buffer, 
                         download_name="question_paper_prediction.txt",
                         as_attachment=True,
                         mimetype='text/plain')
    
    elif format == 'docx':
        doc = Document()
        doc.add_heading('Question Paper Prediction', 0)
        
        # Split by sections and add to document
        sections = ["SECTION 1: Previously Asked Repeated Questions", 
                   "SECTION 2: Most Probable Questions", 
                   "SECTION 3: Generated Question Paper"]
        
        current_text = full_response
        for section in sections:
            if section in current_text:
                section_start = current_text.find(section)
                doc.add_heading(section, level=1)
                
                # Find the next section or end of text
                next_section_start = len(current_text)
                for next_section in sections:
                    if next_section != section and next_section in current_text[section_start+len(section):]:
                        next_start = current_text.find(next_section, section_start+len(section))
                        if next_start < next_section_start:
                            next_section_start = next_start
                
                section_content = current_text[section_start+len(section):next_section_start].strip()
                
                # Add paragraphs
                for para in section_content.split('\n\n'):
                    if para.strip():
                        doc.add_paragraph(para.strip())
        
        # Save to temporary file then serve
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        return send_file(temp_file.name, 
                         download_name="question_paper_prediction.docx", 
                         as_attachment=True)
    
    else:
        flash('Invalid download format.')
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)